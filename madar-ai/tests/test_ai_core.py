import io
import sys
import types

import pytest
from docx import Document

from models.cv_parser import CVParser
from models.embeddings import cosine_similarity
from models.matcher import JobStudentMatcher
from models.skill_extractor import SkillExtractor
from models.job_analyzer import JobAnalyzer
from services.learning_resources import resources_for_skill
from utils.text_cleaner import clean_text


def test_clean_text_normalizes_whitespace_and_keeps_arabic():
    cleaned = clean_text("  مهندس   برمجيات\r\nPython\tReact  ")
    assert "مهندس برمجيات" in cleaned
    assert "python react" in cleaned


def test_docx_text_extraction_reads_paragraphs_and_tables():
    document = Document()
    document.add_paragraph("Madar Candidate")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python"
    stream = io.BytesIO()
    document.save(stream)

    text = CVParser()._extract_text_from_docx(stream.getvalue())

    assert "Madar Candidate" in text
    assert "Python" in text


def test_docx_cv_parsing_extracts_structured_profile_data(monkeypatch):
    document = Document()
    document.add_heading("Madar CV Candidate", 0)
    document.add_paragraph("Backend Engineer")
    document.add_paragraph("candidate@example.test")
    document.add_paragraph("+967711111111")
    document.add_heading("Skills", level=1)
    document.add_paragraph("Python, FastAPI, MongoDB")
    document.add_heading("Experience", level=1)
    document.add_paragraph("Backend Engineer\nMadar Systems\n2022 - 2025\nBuilt FastAPI services.")
    document.add_heading("Education", level=1)
    document.add_paragraph("Bachelor of Software Engineering\nSana'a University\n2024")
    document.add_heading("Projects", level=1)
    document.add_paragraph("- Student matching platform using Python and MongoDB")
    stream = io.BytesIO()
    document.save(stream)
    monkeypatch.setattr("models.cv_parser.generate_embedding", lambda _text: [0.1, 0.2, 0.3])

    parsed = CVParser().parse(stream.getvalue(), "candidate.docx")

    assert parsed.personal_info.name == "Madar CV Candidate"
    assert parsed.personal_info.email == "candidate@example.test"
    assert parsed.personal_info.phone == "+967711111111"
    assert {"Python", "FastAPI", "MongoDB"}.issubset({skill["name"] for skill in parsed.skills})
    assert parsed.experience
    assert parsed.education
    assert parsed.projects
    assert parsed.embedding == [0.1, 0.2, 0.3]
    assert parsed.confidence > 0


def test_pdf_text_extraction_reads_all_pages(monkeypatch):
    class Page:
        def __init__(self, text):
            self.text = text

        def extract_text(self):
            return self.text

    fake_module = types.SimpleNamespace(PdfReader=lambda _: types.SimpleNamespace(pages=[Page("First"), Page("Second")]))
    monkeypatch.setitem(sys.modules, "PyPDF2", fake_module)

    assert CVParser()._extract_text_from_pdf(b"%PDF-test") == "First\nSecond"


def test_cv_parser_rejects_documents_without_readable_text(monkeypatch):
    parser = CVParser()
    monkeypatch.setattr(parser, "_extract_text_from_pdf", lambda _: "")

    with pytest.raises(ValueError, match="No readable text"):
        parser.parse(b"%PDF-empty", "empty.pdf")


def test_skill_extraction_deduplicates_javascript_aliases_and_arabic():
    skills = SkillExtractor().extract_skills("Advanced JS and JavaScript. خبرة في بايثون وحل المشكلات.")
    names = [item.name for item in skills]

    assert names.count("JavaScript") == 1
    assert "Python" in names
    assert "Problem Solving" in names


def test_cosine_similarity_validates_dimensions_and_values():
    assert cosine_similarity([1.0, 0.0], [1.0, 0.0]) == pytest.approx(1.0)
    assert cosine_similarity([1.0, 0.0], [0.0, 1.0]) == pytest.approx(0.0)
    with pytest.raises(ValueError):
        cosine_similarity([1.0], [1.0, 2.0])


def test_match_score_applies_mandatory_skill_penalty():
    matcher = JobStudentMatcher()
    common = dict(
        student_skills=[{"name": "Python", "level": 0.9}],
        student_embedding=[1.0, 0.0],
        job_embedding=[1.0, 0.0],
        student_projects=["Python service"],
        job_projects_hint=["Python"],
    )
    without_missing = matcher.calculate_match(
        **common,
        job_required_skills=[{"name": "Python", "weight": 1, "required": True, "requiredLevel": 0.7}],
    )
    with_missing = matcher.calculate_match(
        **common,
        job_required_skills=[
            {"name": "Python", "weight": 1, "required": True, "requiredLevel": 0.7},
            {"name": "Kubernetes", "weight": 1, "required": True, "requiredLevel": 0.6},
        ],
    )

    assert 0 <= with_missing["overallScore"] <= 100
    assert with_missing["mandatorySkillsPenalty"] > 0
    assert with_missing["overallScore"] < without_missing["overallScore"]


def test_acceptance_probability_is_labeled_as_heuristic_without_history():
    result = JobStudentMatcher().calculate_acceptance_probability(
        match_score=75,
        student_projects_count=2,
        student_certifications_count=1,
        historical_sample_size=0,
    )

    assert 0 <= result["score"] <= 100
    assert result["method"] == "heuristic_estimate"
    assert result["historicalSampleSize"] == 0


@pytest.mark.parametrize("text, expected", [
    ("Three years of experience required", 3),
    ("خبرة لا تقل عن ثلاث سنوات", 3),
    ("خبرة ٤ سنوات", 4),
])
def test_job_experience_understands_written_and_arabic_numbers(text, expected):
    assert JobAnalyzer()._minimum_experience(text) == expected


def test_learning_resources_use_verified_https_catalog_without_fabrication():
    supported_skills = [
        "Machine Learning",
        "REST API",
        "CI/CD",
        "Project Management",
        "Process Engineering",
        "Safety",
    ]

    for skill in supported_skills:
        resources = resources_for_skill(skill)
        assert resources
        assert all(resource["url"].startswith("https://") for resource in resources)
        assert all(resource["skillName"] == skill for resource in resources)
        assert all(resource["reason"] for resource in resources)

    react_resources = resources_for_skill("React")
    python_resources = resources_for_skill("Python")
    assert any(resource["provider"] == "Coursera / Meta" for resource in react_resources)
    assert any(resource["provider"] == "edX / HarvardX" for resource in python_resources)

    assert resources_for_skill("unsupported-private-skill") == []
