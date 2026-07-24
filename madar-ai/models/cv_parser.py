"""
MADAR AI Engine - CV/Resume Parser

Parses PDF and DOCX files to extract structured information including
personal details, skills, experience, education, and projects.
"""

import io
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

from models.embeddings import generate_embedding
from models.skill_extractor import ExtractedSkill, SkillExtractor
from utils.logger import get_logger
from utils.text_cleaner import clean_text

logger = get_logger(__name__)


@dataclass
class PersonalInfo:
    """Personal information extracted from a CV."""

    name: str = ""
    professional_title: str = ""
    email: str = ""
    phone: str = ""
    whatsapp: str = ""
    location: str = ""
    city: str = ""
    country: str = ""
    address: str = ""
    linkedin: str = ""
    github: str = ""
    portfolio: str = ""
    website: str = ""


@dataclass
class ExperienceEntry:
    """A single work experience entry."""

    title: str = ""
    company: str = ""
    years: float = 0.0
    duration_text: str = ""
    description: str = ""
    location: str = ""


@dataclass
class EducationEntry:
    """A single education entry."""

    degree: str = ""
    institution: str = ""
    year: int = 0
    field_of_study: str = ""
    gpa: str = ""


@dataclass
class ParsedCV:
    """Complete parsed CV data structure."""

    personal_info: PersonalInfo = field(default_factory=PersonalInfo)
    skills: List[Dict[str, Any]] = field(default_factory=list)
    experience: List[Dict[str, Any]] = field(default_factory=list)
    education: List[Dict[str, Any]] = field(default_factory=list)
    projects: List[Dict[str, Any]] = field(default_factory=list)
    certifications: List[Dict[str, Any]] = field(default_factory=list)
    courses: List[Dict[str, Any]] = field(default_factory=list)
    languages: List[str] = field(default_factory=list)
    soft_skills: List[str] = field(default_factory=list)
    tools: List[str] = field(default_factory=list)
    volunteer_work: List[Dict[str, Any]] = field(default_factory=list)
    awards: List[str] = field(default_factory=list)
    achievements: List[str] = field(default_factory=list)
    publications: List[str] = field(default_factory=list)
    references: List[str] = field(default_factory=list)
    additional_sections: Dict[str, List[str]] = field(default_factory=dict)
    summary: str = ""
    strengths: List[str] = field(default_factory=list)
    weaknesses: List[str] = field(default_factory=list)
    suggestions: List[str] = field(default_factory=list)
    confidence: float = 0.0
    embedding: Optional[List[float]] = field(default=None, repr=False)
    raw_text: str = ""

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        result = {
            "personalInfo": {
                "name": self.personal_info.name,
                "professionalTitle": self.personal_info.professional_title,
                "email": self.personal_info.email,
                "phone": self.personal_info.phone,
                "whatsapp": self.personal_info.whatsapp,
                "location": self.personal_info.location,
                "city": self.personal_info.city,
                "country": self.personal_info.country,
                "address": self.personal_info.address,
                "linkedin": self.personal_info.linkedin,
                "github": self.personal_info.github,
                "portfolio": self.personal_info.portfolio,
                "website": self.personal_info.website,
            },
            "skills": self.skills,
            "softSkills": self.soft_skills,
            "tools": self.tools,
            "experience": self.experience,
            "education": self.education,
            "projects": self.projects,
            "certifications": self.certifications,
            "courses": self.courses,
            "languages": self.languages,
            "volunteerWork": self.volunteer_work,
            "awards": self.awards,
            "achievements": self.achievements,
            "publications": self.publications,
            "references": self.references,
            "additionalSections": self.additional_sections,
            "summary": self.summary,
            "raw_text": self.raw_text,
            "strengths": self.strengths,
            "weaknesses": self.weaknesses,
            "suggestions": self.suggestions,
            "confidence": self.confidence,
        }
        if self.embedding is not None:
            result["embedding"] = self.embedding
        return result


class CVParser:
    """Parser for CV/resume documents (PDF and DOCX).

    Extracts structured information from uploaded CV files including
    personal details, skills, work experience, education, and projects.
    """

    def __init__(self):
        """Initialize the CV parser with skill extractor."""
        self.skill_extractor = SkillExtractor()

    def parse(self, file_content: bytes, filename: str) -> ParsedCV:
        """Parse a CV file and extract structured information.

        Args:
            file_content: Raw bytes of the uploaded file.
            filename: Original filename (used to determine file type).

        Returns:
            ParsedCV: Structured CV data.

        Raises:
            ValueError: If the file type is not supported.
        """
        # Extract text based on file type
        if filename.lower().endswith(".pdf"):
            raw_text = self._extract_text_from_pdf(file_content)
        elif filename.lower().endswith(".docx"):
            raw_text = self._extract_text_from_docx(file_content)
        else:
            raise ValueError(
                f"Unsupported file type: {filename}. "
                "Only PDF and DOCX files are supported."
            )

        if not raw_text or not raw_text.strip():
            logger.warning("No text extracted from CV", file_name=filename)
            raise ValueError("No readable text could be extracted from the CV")

        logger.info(
            "Extracted text from CV",
            file_name=filename,
            text_length=len(raw_text),
        )

        # Parse structured information
        parsed = self._parse_text(raw_text)
        parsed.raw_text = raw_text

        # Generate AI analysis: strengths, weaknesses, suggestions, confidence
        try:
            parsed.strengths = self._generate_strengths(parsed)
            parsed.weaknesses = self._generate_weaknesses(parsed)
            parsed.suggestions = self._generate_suggestions(parsed)
            parsed.confidence = self._calculate_confidence(parsed)
        except Exception as e:
            logger.error("Failed to generate AI analysis", error=str(e))

        # Generate embedding for the full CV
        try:
            parsed.embedding = generate_embedding(raw_text[:3000])
        except Exception as e:
            logger.error("Failed to generate CV embedding", error=str(e))

        return parsed

    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from a PDF file.

        Args:
            file_content: Raw PDF bytes.

        Returns:
            str: Extracted text content.
        """
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(file_content))
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            extracted = "\n".join(text_parts)
            if extracted.strip():
                return extracted

            return self._extract_text_from_scanned_pdf(file_content)
        except Exception as e:
            logger.error("PDF text extraction failed", error=str(e))
            return ""

    def _extract_text_from_scanned_pdf(self, file_content: bytes) -> str:
        """Extract text from scanned PDFs when OCR dependencies are available."""
        try:
            from pdf2image import convert_from_bytes
            import pytesseract

            pages = convert_from_bytes(file_content, dpi=220)
            text_parts = []
            for page in pages[:10]:
                page_text = pytesseract.image_to_string(page, lang="ara+eng")
                if page_text.strip():
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.warning(
                "Scanned PDF OCR unavailable or failed",
                error=str(e),
            )
            return ""

    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from a DOCX file.

        Args:
            file_content: Raw DOCX bytes.

        Returns:
            str: Extracted text content.
        """
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n".join(text_parts)
        except Exception as e:
            logger.error("DOCX text extraction failed", error=str(e))
            return ""

    def _parse_text(self, text: str) -> ParsedCV:
        """Parse structured information from CV text.

        Args:
            text: Cleaned CV text content.

        Returns:
            ParsedCV: Structured CV data.
        """
        parsed = ParsedCV()

        # Extract personal information
        parsed.personal_info = self._extract_personal_info(text)

        # Extract skills
        skills = self.skill_extractor.extract_skills(text)
        parsed.skills = [
            {
                "name": s.name,
                "category": s.category,
                "confidence": s.confidence,
            }
            for s in skills
        ]

        # Extract work experience
        parsed.experience = self._extract_experience(text)

        # Extract education
        parsed.education = self._extract_education(text)

        # Extract projects
        parsed.projects = self._extract_projects(text)

        # Extract certifications
        parsed.certifications = self._extract_certifications(text)

        parsed.courses = self._extract_named_entries(
            text,
            ["courses", "training", "workshops", "الدورات", "التدريب", "ورش العمل"],
        )
        parsed.languages = self._extract_languages(text)
        parsed.soft_skills = self._extract_soft_skills(text)
        parsed.tools = self._extract_tools(parsed.skills)
        parsed.volunteer_work = self._extract_structured_section(
            text,
            ["volunteer", "volunteering", "community", "التطوع", "العمل التطوعي"],
            ["title", "description"],
        )
        parsed.awards = self._extract_named_entries(
            text,
            ["awards", "honors", "الجوائز", "التكريم"],
        )
        parsed.achievements = self._extract_named_entries(
            text,
            ["achievements", "accomplishments", "الإنجازات", "الانجازات"],
        )
        parsed.publications = self._extract_named_entries(
            text,
            ["publications", "research", "papers", "المنشورات", "الأبحاث", "الابحاث"],
        )
        parsed.references = [
            item["name"]
            for item in self._extract_named_entries(
            text,
            ["references", "المراجع"],
            )
            if item.get("name")
        ]
        parsed.additional_sections = self._extract_additional_sections(text)
        self._apply_arabic_section_fallbacks(parsed, text)

        # Generate summary
        parsed.summary = self._generate_summary(parsed)

        return parsed

    def _extract_personal_info(self, text: str) -> PersonalInfo:
        """Extract personal information from CV text.

        Args:
            text: CV text content.

        Returns:
            PersonalInfo: Extracted personal details.
        """
        info = PersonalInfo()

        email_match = re.search(
            r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
            text,
        )
        if email_match:
            info.email = email_match.group(0)

        phone_patterns = [
            re.compile(r"(?:\+?966|0)?\s?5\d{8}"),
            re.compile(r"\+?[1-9][\d\s().-]{7,18}"),
            re.compile(r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"),
        ]
        for pattern in phone_patterns:
            phone_match = pattern.search(text)
            if phone_match:
                info.phone = self._clean_inline_value(phone_match.group(0))
                break

        whatsapp_match = re.search(
            r"(?:whatsapp|واتساب|واتس)\s*[:\-]?\s*(\+?[0-9][0-9\s().-]{7,18})",
            text,
            re.IGNORECASE,
        )
        if whatsapp_match:
            info.whatsapp = self._clean_inline_value(whatsapp_match.group(1))

        linkedin_match = re.search(
            r"(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9_\-/%]+|linkedin:\s*[A-Za-z0-9_-]+",
            text,
            re.IGNORECASE,
        )
        if linkedin_match:
            info.linkedin = self._normalize_url(
                linkedin_match.group(0).replace("linkedin:", "linkedin.com/in/")
            )

        github_match = re.search(
            r"(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_.\-/%]+|github:\s*[A-Za-z0-9_-]+",
            text,
            re.IGNORECASE,
        )
        if github_match:
            info.github = self._normalize_url(
                github_match.group(0).replace("github:", "github.com/")
            )

        urls = re.findall(
            r"(?:https?://)?(?:www\.)?[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/[^\s,;]*)?",
            text,
        )
        for url in urls:
            normalized_url = self._normalize_url(url)
            lower_url = normalized_url.lower()
            if "linkedin.com" in lower_url or "github.com" in lower_url:
                continue
            if not info.website:
                info.website = normalized_url
            if not info.portfolio and any(
                token in lower_url
                for token in ["portfolio", "behance", "dribbble", "notion", "vercel", "netlify"]
            ):
                info.portfolio = normalized_url

        lines = [line.strip() for line in text.strip().split("\n") if line.strip()]
        ignored_name_tokens = {
            "email", "phone", "mobile", "address", "linkedin", "github", "summary",
            "objective", "experience", "education", "skills", "portfolio", "website",
            "البريد", "الهاتف", "الجوال", "العنوان", "الملخص", "الخبرات", "التعليم", "المهارات",
        }
        for line in lines[:8]:
            lower_line = line.lower()
            if (
                2 < len(line) < 70
                and "@" not in line
                and "http" not in lower_line
                and not re.search(r"\d", line)
                and not any(token in lower_line for token in ignored_name_tokens)
            ):
                words = line.split()
                if 2 <= len(words) <= 5:
                    info.name = line
                    break

        if info.name:
            for line in lines[:10]:
                lower_line = line.lower()
                if (
                    line != info.name
                    and 3 <= len(line) <= 120
                    and "@" not in line
                    and "http" not in lower_line
                    and not re.search(r"\d", line)
                    and not any(token in lower_line for token in ignored_name_tokens)
                ):
                    info.professional_title = line
                    break

        location_match = re.search(
            r"(?:Location|Address|City|Country|المكان|العنوان|المدينة|الدولة)\s*[:\-]?\s*([^\n]+)",
            text,
            re.IGNORECASE,
        )
        if location_match:
            info.location = self._clean_inline_value(location_match.group(1))
        else:
            cities = [
                "Riyadh", "Jeddah", "Dammam", "Khobar", "Mecca", "Medina",
                "Abha", "Tabuk", "Hail", "Taif", "Buraidah",
                "الرياض", "جدة", "الدمام", "الخبر", "مكة", "المدينة",
                "أبها", "تبوك", "حائل", "الطائف", "بريدة",
            ]
            for city in cities:
                if re.search(rf"\b{re.escape(city)}\b", text, re.IGNORECASE):
                    info.location = city
                    break

        if info.location:
            parts = [p.strip() for p in re.split(r"[,،|-]", info.location) if p.strip()]
            if parts:
                info.city = parts[0]
            if len(parts) > 1:
                info.country = parts[-1]
            info.address = info.location

        return info

    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience entries from CV text.

        Args:
            text: CV text content.

        Returns:
            List of experience entry dictionaries.
        """
        experience = []

        # Find experience section
        exp_section = self._extract_section(
            text,
            [
                r"experience",
                r"work experience",
                r"professional experience",
                r"employment",
                r"career",
                r"الخبرات",
                r"الخبرة العملية",
                r"العمل",
            ],
        )

        if not exp_section:
            return experience

        # Split into individual entries (heuristic: split by dates or company patterns)
        entries = self._split_experience_entries(exp_section)

        for entry_text in entries[:10]:  # Limit to 10 entries
            exp_entry = self._parse_experience_entry(entry_text)
            if exp_entry.title or exp_entry.company:
                experience.append(
                    {
                        "title": exp_entry.title,
                        "company": exp_entry.company,
                        "years": exp_entry.years,
                        "duration": exp_entry.duration_text,
                        "description": exp_entry.description[:200]
                        if exp_entry.description
                        else "",
                        "location": exp_entry.location,
                    }
                )

        return experience

    def _split_experience_entries(self, section_text: str) -> List[str]:
        """Split experience section into individual entries.

        Uses date patterns and bullet points as separators.
        """
        # Split by common date patterns
        date_pattern = re.compile(
            r"(?:^|\n)\s*(?:\d{1,2}/\d{2,4}|\d{4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*\d{4}|Present|Current|الحالي)"
        )

        parts = date_pattern.split(section_text)
        entries = [p.strip() for p in parts if len(p.strip()) > 10]

        if not entries:
            # Fallback: split by bullet points or double newlines
            entries = [
                e.strip()
                for e in re.split(r"\n[•\-\*]\s|\n\n", section_text)
                if len(e.strip()) > 10
            ]

        return entries

    def _parse_experience_entry(self, entry_text: str) -> ExperienceEntry:
        """Parse a single experience entry."""
        entry = ExperienceEntry()

        lines = entry_text.strip().split("\n")
        if not lines:
            return entry

        # First non-empty line is usually the title
        for line in lines:
            line = line.strip()
            if line and len(line) > 2:
                if not entry.title:
                    entry.title = line[:100]
                elif not entry.company:
                    entry.company = line[:100]
                    break

        # Extract duration/years
        duration_pattern = re.compile(
            r"(\d{4})\s*(?:–|-|to|حتى)\s*(\d{4}|Present|Current|الحالي)"
        )
        duration_match = duration_pattern.search(entry_text)
        if duration_match:
            start_year = int(duration_match.group(1))
            end_str = duration_match.group(2)
            if end_str.lower() in ("present", "current", "الحالي"):
                end_year = 2024
            else:
                end_year = int(end_str)
            entry.years = max(0, end_year - start_year)
            entry.duration_text = duration_match.group(0)

        # Extract description (remaining lines)
        description_lines = []
        for line in lines[2:]:
            line = line.strip()
            if line and line not in (entry.title, entry.company):
                description_lines.append(line)
        entry.description = " ".join(description_lines)

        return entry

    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education entries from CV text.

        Args:
            text: CV text content.

        Returns:
            List of education entry dictionaries.
        """
        education = []

        edu_section = self._extract_section(
            text,
            [
                r"education",
                r"academic",
                r"qualifications",
                r"التعليم",
                r"المؤهلات",
                r"الشهادات",
            ],
        )

        if not edu_section:
            return education

        # Split entries
        entries = self._split_education_entries(edu_section)

        for entry_text in entries[:5]:
            edu_entry = self._parse_education_entry(entry_text)
            if edu_entry.degree or edu_entry.institution:
                education.append(
                    {
                        "degree": edu_entry.degree,
                        "institution": edu_entry.institution,
                        "year": edu_entry.year,
                        "fieldOfStudy": edu_entry.field_of_study,
                        "gpa": edu_entry.gpa,
                    }
                )

        return education

    def _split_education_entries(self, section_text: str) -> List[str]:
        """Split education section into individual entries."""
        # Split by degree patterns or bullet points
        entries = [
            e.strip()
            for e in re.split(r"\n[•\-\*]\s|\n\n", section_text)
            if len(e.strip()) > 5
        ]
        return entries if entries else [section_text]

    def _parse_education_entry(self, entry_text: str) -> EducationEntry:
        """Parse a single education entry."""
        entry = EducationEntry()

        lines = entry_text.strip().split("\n")

        # Extract degree
        degree_keywords = [
            "bachelor",
            "master",
            "phd",
            "doctorate",
            "associate",
            "diploma",
            "certificate",
            "licence",
            "bsc",
            "msc",
            "ba",
            "ma",
            "bs",
            "ms",
            "meng",
            "beng",
            "mba",
            "بكالوريوس",
            "ماجستير",
            "دكتوراه",
            "دبلوم",
            "شهادة",
        ]

        text_lower = entry_text.lower()
        for keyword in degree_keywords:
            if keyword in text_lower:
                # Extract the full degree line
                for line in lines:
                    if keyword in line.lower():
                        entry.degree = line.strip()[:150]
                        break
                break

        # Extract year
        year_pattern = re.compile(r"\b(19\d{2}|20\d{2})\b")
        year_match = year_pattern.search(entry_text)
        if year_match:
            year = int(year_match.group(1))
            if 1980 <= year <= 2030:
                entry.year = year

        # Extract institution (university name)
        uni_pattern = re.compile(
            r"(?:University|College|Institute|School|Academy|جامعة|كلية|معهد)[\s\w]*",
            re.IGNORECASE,
        )
        uni_match = uni_pattern.search(entry_text)
        if uni_match:
            entry.institution = uni_match.group(0).strip()[:150]

        # Extract field of study
        field_pattern = re.compile(
            r"(?:in|of|major|specialization|تخصص)\s+([\w\s]+(?:Engineering|Science|Computer|IT|Business|Management|Arts|الحاسب|الهندسة|العلوم|الأعمال|الإدارة))",
            re.IGNORECASE,
        )
        field_match = field_pattern.search(entry_text)
        if field_match:
            entry.field_of_study = field_match.group(1).strip()[:100]

        # Extract GPA
        gpa_pattern = re.compile(r"(?:GPA|gpa)[\s:]*(\d+\.?\d*)")
        gpa_match = gpa_pattern.search(entry_text)
        if gpa_match:
            entry.gpa = gpa_match.group(1)

        return entry

    def _extract_projects(self, text: str) -> List[Dict[str, Any]]:
        """Extract project descriptions from CV text."""
        projects = []

        proj_section = self._extract_section(
            text,
            [
                r"projects",
                r"personal projects",
                r"side projects",
                r"academic projects",
                r"المشاريع",
                r"مشاريع",
            ],
        )

        if not proj_section:
            return projects

        # Split by bullet points
        entries = [
            e.strip()
            for e in re.split(r"\n[•\-\*]\s|\n\n", proj_section)
            if len(e.strip()) > 5
        ]

        for entry in entries[:8]:
            # Take first line as project name
            first_line = entry.split("\n")[0].strip()
            if first_line and len(first_line) < 200:
                technologies = [
                    skill.name
                    for skill in self.skill_extractor.extract_skills(entry)[:12]
                ]
                projects.append(
                    {
                        "name": first_line,
                        "title": first_line,
                        "description": self._compact_text(entry),
                        "technologies": technologies,
                        "role": self._extract_labeled_value(entry, ["role", "الدور"]),
                    }
                )

        return projects

    def _extract_certifications(self, text: str) -> List[Dict[str, Any]]:
        """Extract certification names from CV text."""
        cert_section = self._extract_section(
            text,
            [
                r"certifications",
                r"certificates",
                r"licenses",
                r"الشهادات",
                r"الاعتمادات",
            ],
        )

        if not cert_section:
            # Try extracting known certifications from full text
            return self._extract_known_certifications(text)

        certs = []
        entries = [
            e.strip()
            for e in re.split(r"\n[•\-\*]\s|\n\n", cert_section)
            if len(e.strip()) > 3
        ]

        for entry in entries[:10]:
            first_line = entry.split("\n")[0].strip()
            if first_line and len(first_line) < 200:
                certs.append(
                    {
                        "name": first_line,
                        "issuer": self._extract_labeled_value(entry, ["issuer", "provider", "الجهة", "المصدر"]),
                        "date": self._extract_first_date(entry),
                    }
                )

        return certs

    def _extract_known_certifications(self, text: str) -> List[Dict[str, Any]]:
        """Extract known certification names from full text."""
        text_lower = text.lower()
        certs = []

        known_certs = [
            "aws certified",
            "ccna",
            "comptia",
            "cissp",
            "ceh",
            "pmp",
            "scrum master",
            "itil",
            "google cloud",
            "azure",
            "oracle certified",
            "red hat",
            "tableau",
            "power bi",
            "six sigma",
            "iso 27001",
            "cka",
            "ckad",
            "toefl",
            "ielts",
        ]

        for cert in known_certs:
            if cert in text_lower:
                # Find the full certification name
                pattern = re.compile(
                    rf"{re.escape(cert)}[\w\s]*", re.IGNORECASE
                )
                match = pattern.search(text)
                if match:
                    cert_name = match.group(0).strip()
                    if cert_name and not any(c["name"].lower() == cert_name.lower() for c in certs):
                        certs.append({"name": cert_name[:100], "issuer": "", "date": ""})

        return certs

    def _extract_languages(self, text: str) -> List[str]:
        section = self._extract_section(text, ["languages", "اللغات"])
        candidates = section or text
        known_languages = [
            "Arabic", "English", "French", "Spanish", "German", "Chinese",
            "Japanese", "Korean", "Urdu", "Hindi", "Turkish",
            "العربية", "الإنجليزية", "الانجليزية", "الفرنسية", "الإسبانية",
        ]
        return self._unique_keep_order(
            [language for language in known_languages if re.search(rf"\b{re.escape(language)}\b", candidates, re.IGNORECASE)]
        )

    def _extract_soft_skills(self, text: str) -> List[str]:
        soft_skill_names = [
            "Communication", "Leadership", "Problem Solving", "Teamwork",
            "Critical Thinking", "Time Management", "Adaptability",
            "Creativity", "Collaboration", "Presentation",
            "التواصل", "القيادة", "حل المشكلات", "العمل الجماعي",
            "إدارة الوقت", "التفكير النقدي", "الإبداع",
        ]
        return self._unique_keep_order(
            [skill for skill in soft_skill_names if re.search(re.escape(skill), text, re.IGNORECASE)]
        )

    def _extract_tools(self, skills: List[Dict[str, Any]]) -> List[str]:
        tool_categories = {"tools", "devops", "cloud", "database", "design", "security"}
        return self._unique_keep_order(
            [
                skill.get("name", "")
                for skill in skills
                if str(skill.get("category", "")).lower() in tool_categories
            ]
        )

    def _extract_named_entries(self, text: str, headers: List[str]) -> List[Dict[str, Any]]:
        section = self._extract_section(text, headers)
        if not section:
            return []
        entries = self._split_section_entries(section)
        result = []
        for entry in entries[:20]:
            name = entry.split("\n")[0].strip()
            if not name:
                continue
            result.append(
                {
                    "name": name[:180],
                    "description": self._compact_text(entry),
                    "date": self._extract_first_date(entry),
                    "provider": self._extract_labeled_value(entry, ["provider", "issuer", "platform", "الجهة", "المنصة"]),
                }
            )
        return result

    def _extract_structured_section(
        self,
        text: str,
        headers: List[str],
        fields: List[str],
    ) -> List[Dict[str, Any]]:
        section = self._extract_section(text, headers)
        if not section:
            return []
        result = []
        for entry in self._split_section_entries(section)[:12]:
            title = entry.split("\n")[0].strip()
            item = {field: "" for field in fields}
            item["title"] = title[:160]
            item["description"] = self._compact_text(entry)
            result.append(item)
        return result

    def _extract_additional_sections(self, text: str) -> Dict[str, List[str]]:
        known_headers = {
            "summary", "objective", "experience", "work experience", "professional experience",
            "education", "academic", "skills", "projects", "certifications", "certificates",
            "courses", "training", "languages", "references", "awards", "achievements",
            "publications", "volunteer", "interests",
            "الملخص", "الهدف", "الخبرات", "التعليم", "المهارات", "المشاريع",
            "الشهادات", "الدورات", "اللغات", "المراجع", "الجوائز", "الإنجازات",
        }
        sections: Dict[str, List[str]] = {}
        lines = text.split("\n")
        current_header = ""
        buffer: List[str] = []

        def flush() -> None:
            nonlocal buffer, current_header
            if current_header and buffer:
                entries = self._split_section_entries("\n".join(buffer))
                if entries:
                    sections[current_header] = entries[:20]
            buffer = []

        for line in lines:
            stripped = line.strip().strip(":|-")
            is_header = (
                2 <= len(stripped) <= 45
                and not re.search(r"[.@]", stripped)
                and not re.search(r"\d{3,}", stripped)
                and (stripped.isupper() or stripped.lower() in known_headers or stripped in known_headers)
            )
            if is_header:
                flush()
                if stripped.lower() not in known_headers and stripped not in known_headers:
                    current_header = stripped
                else:
                    current_header = ""
                continue
            if current_header:
                buffer.append(line)
        flush()
        return sections

    def _split_section_entries(self, section_text: str) -> List[str]:
        entries = [
            entry.strip()
            for entry in re.split(r"\n[•\-\*]\s|\n\d+[\.)]\s|\n\n", section_text)
            if len(entry.strip()) > 3
        ]
        return entries if entries else [section_text.strip()]

    def _extract_labeled_value(self, text: str, labels: List[str]) -> str:
        for label in labels:
            match = re.search(rf"{re.escape(label)}\s*[:\-]\s*([^\n]+)", text, re.IGNORECASE)
            if match:
                return self._clean_inline_value(match.group(1))
        return ""

    def _extract_first_date(self, text: str) -> str:
        match = re.search(
            r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})\b",
            text,
            re.IGNORECASE,
        )
        return match.group(0) if match else ""

    def _normalize_url(self, url: str) -> str:
        cleaned = self._clean_inline_value(url)
        if cleaned and not cleaned.startswith(("http://", "https://")):
            return f"https://{cleaned}"
        return cleaned

    def _clean_inline_value(self, value: str) -> str:
        return re.sub(r"\s+", " ", value or "").strip(" \t\r\n|,;")

    def _compact_text(self, value: str) -> str:
        return self._clean_inline_value(value.replace("\n", " "))

    def _unique_keep_order(self, values: List[str]) -> List[str]:
        seen = set()
        result = []
        for value in values:
            cleaned = self._clean_inline_value(value)
            key = cleaned.lower()
            if cleaned and key not in seen:
                seen.add(key)
                result.append(cleaned)
        return result

    def _extract_section(
        self, text: str, section_headers: List[str]
    ) -> str:
        """Extract a section from CV text by its header.

        Args:
            text: Full CV text.
            section_headers: List of possible section header patterns.

        Returns:
            str: Section text content, or empty string if not found.
        """
        text_lower = text.lower()
        lines = text.split("\n")
        lines_lower = text_lower.split("\n")

        start_idx = -1
        for i, line in enumerate(lines_lower):
            stripped = line.strip()
            for header in section_headers:
                # Match header as a standalone line
                if re.search(
                    rf"^\s*{header}\s*[\:\-\|]*\s*$", stripped, re.IGNORECASE
                ):
                    start_idx = i + 1
                    break
            if start_idx >= 0:
                break

        if start_idx < 0:
            return ""

        # Find the end of the section (next section header)
        end_idx = len(lines)
        section_end_patterns = [
            r"^\s*(?:experience|education|skills|projects|certifications|languages|references|summary|objective|interests|activities|publications|awards|الخبرات|التعليم|المهارات|المشاريع|الشهادات|اللغات|المراجع|الملخص|الاهتمامات)\s*[\:\-\|]*\s*$"
        ]

        for i in range(start_idx + 1, len(lines)):
            stripped = lines_lower[i].strip()
            for pattern in section_end_patterns:
                if re.search(pattern, stripped, re.IGNORECASE):
                    end_idx = i
                    break
            if end_idx < len(lines):
                break

        section_text = "\n".join(lines[start_idx:end_idx]).strip()
        return section_text

    def _apply_arabic_section_fallbacks(self, parsed: ParsedCV, text: str) -> None:
        """Fill empty entities using correctly encoded Arabic section labels."""
        if not parsed.courses:
            parsed.courses = self._extract_named_entries(text, ["الدورات", "التدريب", "ورش العمل"])
        if not parsed.languages:
            language_section = self._extract_section(text, ["اللغات"])
            names = ["العربية", "الإنجليزية", "الانجليزية", "الفرنسية", "الإسبانية"]
            parsed.languages = [name for name in names if name in language_section]
        if not parsed.volunteer_work:
            parsed.volunteer_work = self._extract_structured_section(text, ["التطوع", "العمل التطوعي"], ["title", "description"])
        if not parsed.awards:
            parsed.awards = [item["name"] for item in self._extract_named_entries(text, ["الجوائز", "التكريم"])]
        if not parsed.achievements:
            parsed.achievements = [item["name"] for item in self._extract_named_entries(text, ["الإنجازات", "الانجازات"])]
        if not parsed.publications:
            parsed.publications = [item["name"] for item in self._extract_named_entries(text, ["المنشورات", "الأبحاث", "الابحاث"])]
        correct_soft_skills = ["التواصل", "القيادة", "حل المشكلات", "العمل الجماعي", "إدارة الوقت", "ادارة الوقت", "التفكير النقدي", "الإبداع"]
        parsed.soft_skills = self._unique_keep_order([*parsed.soft_skills, *[name for name in correct_soft_skills if name in text]])

    def _generate_summary(self, parsed: ParsedCV) -> str:
        """Generate a text summary of the parsed CV.

        Args:
            parsed: The parsed CV data.

        Returns:
            str: A brief summary of the CV.
        """
        parts = []

        if parsed.personal_info.name:
            parts.append(f"Name: {parsed.personal_info.name}")

        if parsed.skills:
            top_skills = [s["name"] for s in parsed.skills[:5]]
            parts.append(f"Top skills: {', '.join(top_skills)}")

        total_years = sum(
            exp.get("years", 0) for exp in parsed.experience
        )
        if total_years > 0:
            parts.append(f"Experience: {total_years:.1f} years")

        if parsed.education:
            degrees = [e.get("degree", "") for e in parsed.education[:2]]
            parts.append(f"Education: {', '.join(d for d in degrees if d)}")

        return "; ".join(parts) if parts else "No summary available"

    def _generate_strengths(self, parsed: ParsedCV) -> List[str]:
        """Generate a list of strengths based on parsed CV data.

        Args:
            parsed: The parsed CV data.

        Returns:
            List of strength descriptions.
        """
        strengths = []

        # Skill-based strengths
        skill_count = len(parsed.skills)
        if skill_count >= 10:
            strengths.append(f"Diverse skill set with {skill_count} identified competencies")
        elif skill_count >= 5:
            strengths.append(f"Solid technical foundation with {skill_count} core skills")

        top_skills = [s["name"] for s in parsed.skills[:3]]
        if top_skills:
            strengths.append(f"Proficiency in key technologies: {', '.join(top_skills)}")

        # Experience-based strengths
        total_years = sum(exp.get("years", 0) for exp in parsed.experience)
        if total_years >= 3:
            strengths.append(f"Substantial work experience ({total_years:.1f} years)")
        elif total_years > 0:
            strengths.append("Practical industry experience")

        if len(parsed.experience) >= 2:
            strengths.append("Exposure to multiple roles and work environments")

        # Education-based strengths
        if parsed.education:
            degrees = [e.get("degree", "") for e in parsed.education if e.get("degree")]
            if degrees:
                strengths.append(f"Formal education: {degrees[0]}")

        # Project-based strengths
        if len(parsed.projects) >= 3:
            strengths.append(f"Strong portfolio with {len(parsed.projects)} projects")
        elif parsed.projects:
            strengths.append("Hands-on project experience")

        # Certification-based strengths
        if parsed.certifications:
            cert_count = len(parsed.certifications)
            if cert_count >= 2:
                strengths.append(f"Committed to professional development ({cert_count} certifications)")

        return strengths if strengths else ["CV submitted for analysis"]

    def _generate_weaknesses(self, parsed: ParsedCV) -> List[str]:
        """Generate a list of weaknesses / improvement areas based on parsed CV data.

        Args:
            parsed: The parsed CV data.

        Returns:
            List of weakness descriptions.
        """
        weaknesses = []

        # Check for missing contact info
        if not parsed.personal_info.email:
            weaknesses.append("No email contact found in CV")
        if not parsed.personal_info.phone:
            weaknesses.append("No phone number listed")

        # Check for thin experience
        if not parsed.experience:
            weaknesses.append("No work experience section detected")
        elif len(parsed.experience) == 1:
            weaknesses.append("Limited work history (single position)")

        total_years = sum(exp.get("years", 0) for exp in parsed.experience)
        if 0 < total_years < 1:
            weaknesses.append("Very limited professional experience (< 1 year)")

        # Check for thin skills
        if len(parsed.skills) < 5:
            weaknesses.append("Skill set appears limited; consider adding more technologies")

        # Check for missing sections
        if not parsed.projects:
            weaknesses.append("No projects section; adding projects strengthens candidacy")

        if not parsed.education:
            weaknesses.append("No education details found")

        if not parsed.certifications:
            weaknesses.append("No professional certifications listed")

        return weaknesses if weaknesses else ["No major weaknesses identified"]

    def _generate_suggestions(self, parsed: ParsedCV) -> List[str]:
        """Generate improvement suggestions based on parsed CV data.

        Args:
            parsed: The parsed CV data.

        Returns:
            List of actionable suggestions.
        """
        suggestions = []

        # Skill-related suggestions
        skill_names = [s["name"].lower() for s in parsed.skills]
        in_demand_skills = [
            "docker", "kubernetes", "aws", "azure", "gcp",
            "react", "typescript", "node.js", "python",
            "sql", "nosql", "mongodb", "postgresql",
            "ci/cd", "git", "github actions", "jenkins",
            "machine learning", "ai", "data analysis",
            "agile", "scrum", "jira",
        ]
        missing_in_demand = [
            s for s in in_demand_skills
            if s not in skill_names and len(skill_names) < 20
        ]
        if missing_in_demand:
            top_missing = missing_in_demand[:5]
            suggestions.append(
                f"Consider learning in-demand skills: {', '.join(top_missing)}"
            )

        # Section improvement suggestions
        if not parsed.personal_info.linkedin:
            suggestions.append("Add a LinkedIn profile URL for professional visibility")
        if not parsed.personal_info.github:
            suggestions.append("Include a GitHub portfolio to showcase code quality")

        if len(parsed.projects) < 3:
            suggestions.append("Add more projects (aim for at least 3) to demonstrate practical skills")

        if not parsed.certifications:
            suggestions.append("Pursue relevant certifications (e.g., AWS, Azure, Scrum) to stand out")

        # Experience quality suggestions
        for exp in parsed.experience:
            if not exp.get("description") or len(exp.get("description", "")) < 50:
                suggestions.append(
                    f"Add detailed bullet-point achievements for your role at {exp.get('company', 'your company')}"
                )
                break  # Only suggest for the first one

        # General formatting suggestions
        suggestions.append("Use quantifiable achievements (e.g., 'Improved performance by 30%')")
        suggestions.append("Ensure consistent formatting throughout the CV")

        return suggestions if suggestions else ["Keep your CV updated with latest skills and experience"]

    def _calculate_confidence(self, parsed: ParsedCV) -> float:
        """Calculate an overall parsing confidence score (0-1).

        Args:
            parsed: The parsed CV data.

        Returns:
            float: Confidence score between 0 and 1.
        """
        score = 0.0
        total_weights = 0.0

        # Personal info weight: 15%
        total_weights += 0.15
        if parsed.personal_info.name:
            score += 0.075
        if parsed.personal_info.email:
            score += 0.0375
        if parsed.personal_info.phone:
            score += 0.0375

        # Skills weight: 25%
        total_weights += 0.25
        skill_count = len(parsed.skills)
        if skill_count >= 10:
            score += 0.25
        elif skill_count >= 5:
            score += 0.1875
        elif skill_count > 0:
            score += 0.125
        else:
            score += 0.0625

        # Experience weight: 25%
        total_weights += 0.25
        exp_count = len(parsed.experience)
        if exp_count >= 3:
            score += 0.25
        elif exp_count >= 1:
            score += 0.1875
        else:
            score += 0.0625

        # Education weight: 15%
        total_weights += 0.15
        if parsed.education:
            score += 0.15
        else:
            score += 0.05

        # Projects weight: 10%
        total_weights += 0.10
        proj_count = len(parsed.projects)
        if proj_count >= 3:
            score += 0.10
        elif proj_count > 0:
            score += 0.075
        else:
            score += 0.025

        # Certifications weight: 10%
        total_weights += 0.10
        if parsed.certifications:
            score += 0.10
        else:
            score += 0.03

        # Normalize to 0-1 range based on total weights
        if total_weights > 0:
            normalized_score = score / total_weights
        else:
            normalized_score = 0.5

        # Round to 2 decimal places
        return round(min(max(normalized_score, 0.0), 1.0), 2)
